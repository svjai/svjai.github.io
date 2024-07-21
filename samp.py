from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Title
doc.add_heading('Specimen Adequacy in Cervical Cytology', 0)

# Overview
doc.add_heading('Overview', level=1)
doc.add_paragraph(
    "Specimen adequacy is crucial in cervical cytology to ensure that the sample collected is sufficient for accurate diagnosis. "
    "Adequacy is determined by evaluating the presence and quantity of specific cells in the sample. This document provides a detailed formula and guidelines for assessing specimen adequacy using AI detections and auto annotations."
)

# Key Cell Types
doc.add_heading('Key Cell Types', level=1)
doc.add_paragraph(
    "The following cell types are considered for determining specimen adequacy:\n"
    "- Endocervical Cells\n"
    "- Superficial Squamous Cells\n"
    "- Intermediate Squamous Cells\n"
    "- Mild Dysplasia (LSIL)\n"
    "- Moderate Dysplasia (HSIL)\n"
    "- Severe Dysplasia (HSIL)\n"
    "- Squamous Cell Carcinoma\n"
    "- Adenocarcinoma\n"
    "- Atypical Glandular Cells of Undetermined Significance (AGUS)\n"
    "- Atypical Squamous Cells of Undetermined Significance (ASCUS)\n"
    "- Bacterial Infection\n"
    "- Fungal Infection\n"
    "- Protozoan Infection\n"
    "- Metaplastic Squamous Cells\n"
    "- Parabasal Cells\n"
    "- Atypical Squamous Cells - Cannot Exclude HSIL (ASC-H)"
)

# Adequacy Formula
doc.add_heading('Adequacy Formula', level=1)
doc.add_paragraph(
    "To determine specimen adequacy, we will use the counts of key cell types detected in all images from a single slide. "
    "The formula involves setting thresholds for the minimum number of cells detected per image and overall cell count across all images."
)

# Step-by-Step Adequacy Determination
doc.add_heading('Step-by-Step Adequacy Determination', level=1)
doc.add_paragraph(
    "1. Analyze Images: AI analyzes a set number of images (e.g., 1000) at a specified magnification (e.g., 30-40x).\n"
    "2. Count Cells: AI detects and counts the relevant cell types in each image.\n"
    "3. Calculate Totals:\n"
    "    - Calculate Total Cell Count (TCC_cell) for each cell type.\n"
    "    - Calculate Average Cell Count per Image (ACCI_cell): ACCI_cell = Total Cell Count (TCC_cell) / Total Image Count (TIC).\n"
    "4. Check Adequacy for Each Cell Type:\n"
    "    - Determine if the average and total counts meet the specified thresholds."
)

# Formula
doc.add_heading('Formula', level=2)
doc.add_paragraph(
    "Specimen Adequacy =\n"
    "    Adequate if (TCC_cell / TIC) ≥ MCC_cell and TCC_cell ≥ OMCC_cell\n"
    "    Inadequate otherwise"
)

# Thresholds for Each Cell Type
doc.add_heading('Thresholds for Each Cell Type', level=1)
doc.add_paragraph(
    "The following are the thresholds for determining the adequacy of each cell type:\n"
)

cell_types = [
    ("Endocervical Cells", 5, 3000),
    ("Superficial Squamous Cells", 5, 3000),
    ("Intermediate Squamous Cells", 5, 3000),
    ("Mild Dysplasia (LSIL)", 2, 1000),
    ("Moderate Dysplasia (HSIL)", 2, 1000),
    ("Severe Dysplasia (HSIL)", 2, 1000),
    ("Squamous Cell Carcinoma", 1, 500),
    ("Adenocarcinoma", 1, 500),
    ("Atypical Glandular Cells of Undetermined Significance (AGUS)", 1, 500),
    ("Atypical Squamous Cells of Undetermined Significance (ASCUS)", 2, 1000),
    ("Bacterial Infection", 1, 500),
    ("Fungal Infection", 1, 500),
    ("Protozoan Infection", 1, 500),
    ("Metaplastic Squamous Cells", 5, 3000),
    ("Parabasal Cells", 5, 3000),
    ("Atypical Squamous Cells - Cannot Exclude HSIL (ASC-H)", 2, 1000),
]

for cell_type, mcc, omcc in cell_types:
    doc.add_paragraph(
        f"{cell_type}:\n"
        f"  - Minimum Cell Count per Image (MCC): {mcc} cells per image\n"
        f"  - Overall Minimum Cell Count (OMCC): {omcc} cells\n"
    )

# Example Calculation
doc.add_heading('Example Calculation', level=1)

# Thresholds
doc.add_heading('Thresholds', level=2)
doc.add_paragraph(
    "Total Image Count (TIC): 1000 images\n"
)

for cell_type, mcc, omcc in cell_types:
    doc.add_paragraph(
        f"{cell_type}:\n"
        f"  - Minimum Cell Count per Image (MCC): {mcc} cells per image\n"
        f"  - Overall Minimum Cell Count (OMCC): {omcc} cells"
    )

# Data
doc.add_heading('Data', level=2)
doc.add_paragraph(
    "Total Cell Count (TCC)_Endocervical: 3200 cells\n"
    "Total Cell Count (TCC)_Superficial Squamous: 3500 cells\n"
    "Total Cell Count (TCC)_Intermediate Squamous: 4000 cells\n"
    "Total Cell Count (TCC)_Mild Dysplasia (LSIL): 1200 cells\n"
    "Total Cell Count (TCC)_Moderate Dysplasia (HSIL): 1000 cells\n"
    "Total Cell Count (TCC)_Severe Dysplasia (HSIL): 800 cells\n"
    "Total Cell Count (TCC)_Squamous Cell Carcinoma: 500 cells\n"
    "Total Cell Count (TCC)_Adenocarcinoma: 300 cells\n"
    "Total Cell Count (TCC)_AGUS: 200 cells\n"
    "Total Cell Count (TCC)_ASCUS: 1200 cells\n"
    "Total Cell Count (TCC)_Bacterial Infection: 400 cells\n"
    "Total Cell Count (TCC)_Fungal Infection: 300 cells\n"
    "Total Cell Count (TCC)_Protozoan Infection: 200 cells\n"
    "Total Cell Count (TCC)_Metaplastic Squamous: 3500 cells\n"
    "Total Cell Count (TCC)_Parabasal: 3800 cells\n"
    "Total Cell Count (TCC)_ASC-H: 1000 cells"
)

# Calculations
doc.add_heading('Calculations', level=2)
doc.add_paragraph(
    "Average Cell Count per Image (ACCI)_Endocervical: 3200 / 1000 = 3.2 cells per image\n"
    "Average Cell Count per Image (ACCI)_Superficial Squamous: 3500 / 1000 = 3.5 cells per image\n"
    "Average Cell Count per Image (ACCI)_Intermediate Squamous: 4000 / 1000 = 4.0 cells per image\n"
    "Average Cell Count per Image (ACCI)_Mild Dysplasia (LSIL): 1200 / 1000 = 1.2 cells per image\n"
    "Average Cell Count per Image (ACCI)_Moderate Dysplasia (HSIL): 1000 / 1000 = 1.0 cells per image\n"
    "Average Cell Count per Image (ACCI)_Severe Dysplasia (HSIL): 800 / 1000 = 0.8 cells per image\n"
    "Average Cell Count per Image (ACCI)_Squamous Cell Carcinoma: 500 / 1000 = 0.5 cells per image\n"
    "Average Cell Count per Image (ACCI)_Adenocarcinoma: 300 / 1000 = 0.3 cells per image\n"
    "Average Cell Count per Image (ACCI)_AGUS: 200 / 1000 = 0.2 cells per image\n"
    "Average Cell Count per Image (ACCI)_ASCUS: 1200 / 1000 = 1.2 cells per image\n"
    "Average Cell Count per Image (ACCI)_Bacterial Infection: 400 / 1000 = 0.4 cells per image\n"
    "Average Cell Count per Image (ACCI)_Fungal Infection: 300 / 1000 = 0.3 cells per image\n"
    "Average Cell Count per Image (ACCI)_Protozoan Infection: 200 / 1000 = 0.2 cells per image\n"
    "Average Cell Count per Image (ACCI)_Metaplastic Squamous: 3500"
)
